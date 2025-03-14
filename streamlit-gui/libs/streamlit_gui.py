import streamlit as st 
import streamlit_authenticator as stauth 
from datetime import datetime, timezone 
import hashlib 
import yaml 
from typing import Dict, Generator, Tuple 
import dropbox 
import pandas as pd 
from docx import Document
from io import BytesIO
from pathlib import Path 
import threading 

from .ai_gateways.gateway import AICompanyGateway 


class StreamlitGUI: 
    def __init__(self, page_title:str, page_icon:str, ai_company:str, ai_model:str, max_tokens:int, system_message:str, generate_summary_prompt:str, auth_required:bool, interviewer_avatar:str, user_avatar:str, closing_messages:Dict[str, str], dropbox_path:str, interview_instructions:str) -> None: 
        """Set up the object

        Args:
            page_title (str): the title for the page 
            page_icon (str): the icon to use for the page 
            ai_company (str): the name of the AI company to use 
            ai_model (str): the name of the AI model to use 
            max_tokens (int): the max number of output tokens for the model 
            system_message (str): the system prompt for the bot 
            generate_summary_prompt (str): the summary prompt for the bot 
            auth_required (bool): True if logins are enabled and authentication is required, False otherwise 
            interviewer_avatar (str): the string for the interviewer avatar
            user_avatar (str): the string for the user avatar 
            closing_messages (Dict[str, str]): a dict that maps closing code to closing message 
            dropbox_path (str): the path to the dropbox data folder to store the transcripts 
            interview_instructions (str): the instructions to display for the bot 
        """
        # set the global vars
        self.page_title = page_title 
        self.page_icon = page_icon 
        self.ai_company = ai_company 
        self.ai_model = ai_model 
        self.max_tokens = max_tokens 
        self.system_message = system_message 
        self.generate_summary_prompt = generate_summary_prompt
        self.auth_required = auth_required 
        self.interviewer_avatar = interviewer_avatar
        self.user_avatar = user_avatar
        self.closing_messages = closing_messages 
        self.dropbox_path = dropbox_path 
        self.interview_instructions = interview_instructions

        # set up the page 
        st.set_page_config(
            page_title=self.page_title, 
            page_icon=self.page_icon
        )

        if self.auth_required: 
            # set up the authentication 
            stauth_config = yaml.safe_load(st.secrets['STREAMLIT_AUTHENTICATOR_CONFIG'])
            self.authenticator = stauth.Authenticate(credentials=stauth_config['credentials'], auto_hash=False)

        # create some containers for the header (where the title will live) and for the chat (where chat history will live) 
        self.header_container = st.container() 
        self.chat_container = st.container() 


    def setup(self) -> None: 
        """Sets up the components of the page"""
        # set the title of the page 
        with self.header_container: 
            st.title(self.page_title) 
        self.setup_session_vars() 
        self.display_login_page() 
        self.display_instructions_button() 
        self.display_quit_interview_button() 
        self.display_restart_interview_button()
        self.display_generate_summary_button() 
        self.display_user_input() 


    def run(self) -> None: 
        """Main function that runs the whole page"""
        self.setup() 
        self.display_message_history() 
        if st.session_state.interview_status and not st.session_state.first_instructions_shown: 
            self.display_instructions()
        if st.session_state.interview_status: 
            self.stream_initial_message()  
        if st.session_state.interview_status and not st.session_state.first_instructions_shown: 
            st.session_state.first_instructions_shown = True 


    # --------------------------------------------------------------------------
    # frontend 
    # --------------------------------------------------------------------------


    def setup_session_vars(self) -> None: 
        """Sets up the session vars

        Initializes any vars that haven't been initialized yet 
        """
        if not self.auth_required: 
            # since we're not doing logins set these session vars to defaults
            st.session_state.authentication_status = True 
            st.session_state.username = 'testuser'
            st.session_state.show_login_form = False 

        if self.auth_required and 'show_login_form' not in st.session_state:
            # flag for whether to show the login page or not 
            st.session_state.show_login_form = True 

        if "interview_status" not in st.session_state: 
            # flag for whether the interview is active or not 
            st.session_state.interview_status = st.session_state.get('authentication_status', False) 

        if "transcript_history" not in st.session_state: 
            # will store the transcript history of the conversation so far
            st.session_state.transcript_history = [] 
            # will store the messages to display on the screen 
            st.session_state.gui_message_history = [] 

        if 'session_id' not in st.session_state and 'username' in st.session_state: 
            # store the start time of the interview 
            st.session_state.start_time = datetime.now(timezone.utc).timestamp() 

            # create and store the session ID of the interview 
            data = f"{st.session_state.username}+{st.session_state.start_time}"
            st.session_state.session_id = hashlib.sha256(data.encode()).hexdigest() 

        if 'quit_button_hit' not in st.session_state: 
            # flag for whether the quit button has been hit or not 
            st.session_state.quit_button_hit = False 

        if 'first_instructions_shown' not in st.session_state: 
            st.session_state.first_instructions_shown = False 


    def display_login_page(self) -> None: 
        """Display the login page and the log out button after authentication success"""
        if self.auth_required: 
            # only do this if we're using logins 
            if st.session_state.show_login_form:
                # if the session state says show the form, then show it 
                with st.empty(): 
                    try: 
                        # show the login form 
                        self.authenticator.login()
                    except Exception as e: 
                        st.error(e) 

                with st.empty(): 
                    if st.session_state.get('authentication_status') is False: 
                        # password is wrong
                        st.error("Username/password is incorrect") 
                    elif st.session_state.get('authentication_status'):
                        # password is right so change the flag and rerun the script
                        st.session_state.show_login_form = False 
                        st.rerun() 

            if st.session_state.get('authentication_status'): 
                # when authentication is confirmed, show the logout button
                with st.sidebar: 
                    # welcome the user 
                    st.write(f"# Welcome {st.session_state.get('name')}")
                    if not st.session_state.quit_button_hit: 
                        # unless we are in quit button hit status, start the interview
                        st.session_state.interview_status = True 
                    # add logout button that runs self.on_logout when hit 
                    self.authenticator.logout(callback=self.on_logout) 


    def display_instructions_button(self) -> None: 
        if not st.session_state.show_login_form: 
            with st.sidebar: 
                st.markdown("To view the instructions for the assignment and the AI interviewer, click instructions below")
                st.button(
                    label="Instructions", 
                    key='instructions', 
                    help='View instructions', 
                    on_click=self.display_instructions
                )


    @st.dialog("AI Referee Report Guide", width='large')
    def display_instructions(self) -> None: 
        st.markdown(self.interview_instructions)


    def display_quit_interview_button(self) -> None: 
        """Displays the quit interview button"""
        if st.session_state.interview_status: 
            # Add 'Quit' button to the side bar 
            with st.sidebar: 
                st.markdown("To end the interview and receive a summary of your chat, click quit below")
                # add the button and runs self.on_quit_button when hit 
                st.button(
                    label="Quit", 
                    key='quit', 
                    help='End the interview', 
                    on_click=self.on_quit_button
                )


    def display_restart_interview_button(self) -> None: 
        """Displays the restart interview when the quit button is hit"""
        if st.session_state.quit_button_hit and not st.session_state.interview_status and not st.session_state.show_login_form: 
            # Only show the restart button if the quit button has been hit, interview is no longer active, and we are not showing the login page
            with st.sidebar: 
                st.markdown("To restart the interview, click restart below")
                # add the button and runs self.on_restart_button when hit 
                st.button(
                    label="Restart", 
                    key='restart', 
                    help='Restart the interview', 
                    on_click=self.on_restart_button
                )


    def display_generate_summary_button(self) -> None: 
        """Displays the generate summary button"""
        if not st.session_state.show_login_form: 
            # button is always displayed unless we are in the login page to allow people to generate the document at any time 
            with st.sidebar: 
                st.markdown("To generate a summary document of the interview, click generate below") 
                # add the button and runs self.on_generate_summary_button when hit 
                st.button(
                    label="Generate", 
                    key='Generate', 
                    help='Generate summary of the interview', 
                    on_click=self.on_generate_summary_button
                )


    def display_message_history(self) -> None: 
        """Displays the full message history so far"""
        if not st.session_state.show_login_form: 
            # always show the message history unless we are showing the login page 
            with self.chat_container: 
                for message in st.session_state.gui_message_history: 
                    # first set the avatar 
                    if message['role'] == 'assistant': 
                        avatar = self.interviewer_avatar 
                    elif message['role'] == 'user': 
                        avatar = self.user_avatar 

                    # now display the message 
                    with st.chat_message(message['role'], avatar=avatar): 
                        st.markdown(message['content']) 


    def display_user_input(self) -> None: 
        """Display the user input """
        if st.session_state.interview_status and not st.session_state.quit_button_hit: 
            # only display the user input section if the interview is active and the quit button hasn't been hit 
            # display the input section and runs self.on_user_input_submit when text submitted 
            st.chat_input(
                placeholder="Your message here", 
                key="user_input", 
                on_submit=self.on_user_input_submit,
            )


    # --------------------------------------------------------------------------
    # backend 
    # --------------------------------------------------------------------------


    def on_logout(self, *args, **kwargs) -> None: 
        """Function that runs when log out button is hit"""
        # stop the interview 
        st.session_state.interview_status = False 
        # show the login form 
        st.session_state.show_login_form = True
        # reset quit button hit 
        st.session_state.quit_button_hit = False
        # reset instructions flag 
        st.session_state.first_instructions_shown = False 

        # remove any other session variable to start over 
        for key in ['transcript_history', 'gui_message_history', 'start_time', 'session_id']: 
            if key in st.session_state:
                del st.session_state[key]


    def on_user_input_submit(self) -> None: 
        """Function that runs when user input is submitted"""
        # get the user inputs 
        text = st.session_state.user_input 
        # display the user input 
        with self.chat_container: 
            with st.chat_message('user', avatar=self.user_avatar): 
                st.markdown(text)

        # save the user input  
        self.save_msg_to_session('user', text)

        # save to the transcript so far to dropbox 
        thread = threading.Thread(target=self.save_transcript_to_dropbox, args=(st.session_state.to_dict(),)) 
        thread.start() 

        # get the response from the AI bot and stream the message 
        client = AICompanyGateway.factory(company=self.ai_company, api_key=st.secrets[f"API_KEY_{self.ai_company.upper()}"]) 
        stream = client.stream_message(model=self.ai_model, messages=st.session_state.gui_message_history.copy(), max_tokens=self.max_tokens, system_message=self.system_message)
        self.stream_message(stream) 


    def on_quit_button(self) -> None: 
        """Function that runs when the quit button is hit"""
        # set the interview state to False as it's over 
        st.session_state.interview_status = False 

        # save a message to the session that the interview is over 
        self.save_msg_to_session('assistant', "You have cancelled the interview.")

        # save transcript to dropbox 
        thread = threading.Thread(target=self.save_transcript_to_dropbox, args=(st.session_state.to_dict(),)) 
        thread.start() 

        # set the quit button to true 
        st.session_state.quit_button_hit = True 


    @st.dialog("Interview Summary Document")
    def on_generate_summary_button(self) -> None: 
        """Function that runs when the generate summary button is hit 

        Creates a pop up dialog that shows a loading spinner and then displays a download button 
        """
        # start the loading spinner and show the time elapsed so far 
        with st.spinner("Generating document", show_time=True):
            # ask the AI to generate a summary 
            client = AICompanyGateway.factory(company=self.ai_company, api_key=st.secrets[f"API_KEY_{self.ai_company.upper()}"]) 
            generate_message = [{'role': 'user', 'content': self.generate_summary_prompt}]
            summary = client.create_message(
                model=self.ai_model, 
                messages=st.session_state.gui_message_history + generate_message, 
                max_tokens=self.max_tokens, 
                system_message=self.system_message 
            )

            # check if there are any closing messages in there 
            _, summary = self.check_closing_messages(summary) 

            # start the word document 
            doc = Document() 
            # add a title 
            doc.add_heading("Interview Summary", 0) 
            # add some details about the report 
            doc.add_paragraph(f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')} by {st.session_state.name}")

            # add section for the summary 
            doc.add_heading("Summary", 1) 
            doc.add_paragraph(summary)

            # add section for the transcript 
            doc.add_heading("Interview Transcript", 1)
            for message in st.session_state.gui_message_history: 
                # first bold the role 
                p = doc.add_paragraph() 
                role_run = p.add_run(f"{message['role'].title()}: ")
                role_run.bold = True 

                # now add the message 
                content_run = p.add_run(message['content']) 

                # italicize the assistant messages 
                if message['role'] == 'assistant': 
                    content_run.italic = True 

            # save the document into BytesIO 
            doc_bytes = BytesIO() 
            doc.save(doc_bytes) 
            doc_bytes.seek(0) 

        # save the document to dropbox 
        thread = threading.Thread(target=self.save_summary_to_dropbox, args=(st.session_state.to_dict(), doc_bytes)) 
        thread.start() 

        # display download button 
        st.markdown("To download the summary document, click download below")
        st.download_button(
            label='Download document', 
            help='Download interview summary document', 
            data=doc_bytes,
            file_name=f"{st.session_state.username}_interview_summary.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            on_click="ignore", 
            icon=":material/download:"
        )


    def on_restart_button(self) -> None: 
        """Function that runs when the restart button is hit"""
        # reset some session state variables 
        for key in ['transcript_history', 'gui_message_history', 'start_time', 'session_id', 'quit_button_hit']: 
            if key in st.session_state:
                del st.session_state[key]
        # restart the interview 
        st.session_state.interview_status = True 


    def stream_initial_message(self) -> None: 
        """Streams the initial message from the AI"""
        if not st.session_state.gui_message_history: 
            # no messages so far, stream initial message 
            client = AICompanyGateway.factory(company=self.ai_company, api_key=st.secrets[f"API_KEY_{self.ai_company.upper()}"]) 
            stream = client.stream_message(model=self.ai_model, messages=[{"role": "user", "content": "Hi!"}], max_tokens=self.max_tokens, system_message=self.system_message)
            self.stream_message(stream) 


    def stream_message(self, stream:Generator) -> None: 
        """Helper function to stream AI messages 

        Args:
            stream (Generator): the generator that contains the messages being streamed 
        """ 
        with self.chat_container: 
            # stream messages within the chat container
            with st.chat_message("assistant", avatar=self.interviewer_avatar): 
                # stream messages as the assistant 
                streamlit_msg = st.empty() # streamlit object for where the message will go 
                msg_so_far = "" # record the message received so far
                for chunk in stream: 
                    # iterate through the stream and add the results 
                    if chunk: 
                        msg_so_far += chunk 
                    found_closing_msg, closing_msg = self.check_closing_messages(msg_so_far) 
                    if found_closing_msg: 
                        streamlit_msg.empty() 
                        break 
                    if len(msg_so_far) > 10: 
                        streamlit_msg.markdown(msg_so_far + "▌")

                # after all the text has streamed
                if found_closing_msg: 
                    # we found a closing message, so display closing message and shut down the conversation 
                    final_msg = closing_msg
                    st.session_state.interview_status = False 
                    st.session_state.quit_button_hit = True 
                else: 
                    # did not find closing message, display the message sent 
                    final_msg = msg_so_far 

                # display the message received 
                streamlit_msg.markdown(final_msg)

                # save the message to the session 
                self.save_msg_to_session('assistant', final_msg)

                # save the transcript to dropbox 
                thread = threading.Thread(target=self.save_transcript_to_dropbox, args=(st.session_state.to_dict(),)) 
                thread.start() 


    # --------------------------------------------------------------------------
    # utils 
    # --------------------------------------------------------------------------


    def save_transcript_to_dropbox(self, session_state:Dict) -> None: 
        """Saves the transcript to dropbox 

        Usually runs in a separate thread to not interrupt the main chatbot experience 

        Args:
            session_state (Dict): the current session state dict to reference inside the thread 
        """
        # creates the path to save to 
        save_fpath = Path(self.dropbox_path)/session_state['username']/f"transcript+{session_state['username']}+{session_state['session_id']}.csv"

        # save the transcript history 
        df = pd.DataFrame(session_state['transcript_history'])
        csv_content = BytesIO() 
        df.to_csv(csv_content, index=False, encoding='utf-8')
        csv_content.seek(0)
        self.save_to_dropbox(csv_content, str(save_fpath))


    def save_summary_to_dropbox(self, session_state:Dict, doc_content:BytesIO) -> None: 
        """Saves the summary docx to dropbox 

        Usually runs in a separate thread to not interrupt the main chatbot experience 

        Args:
            session_state (Dict): the current session state dict to reference inside the thread 
            doc_content (BytesIO): the docx data to save 
        """
        # create the path to save to 
        save_fpath = Path(self.dropbox_path)/session_state['username']/f"summary_document+{session_state['username']}+{session_state['session_id']}+{int(datetime.now(timezone.utc).timestamp())}.docx"

        # save the content to dropbox 
        self.save_to_dropbox(doc_content, str(save_fpath))


    def save_to_dropbox(self, content:BytesIO, save_fpath:str) -> None: 
        """Saves some content to dropbox 

        Usually runs in a separate thread to not interrupt the main chatbot experience 

        Args:
            content (BytesIO): the content to save
            save_fpath (str): the path to save to 
        """
        # create the dropbox client 
        dbx = dropbox.Dropbox(oauth2_refresh_token=st.secrets['REFRESH_TOKEN_DROPBOX'], app_key=st.secrets['APP_KEY_DROPBOX'], app_secret=st.secrets['APP_SECRET_DROPBOX']) 

        # upload the file to dropbox and overwrite the existing file 
        dbx.files_upload(
            content.read(), 
            save_fpath, 
            mode=dropbox.files.WriteMode("overwrite")
        ) 


    def save_msg_to_session(self, role:str, content:str) -> None: 
        """Saves messages in the conversation to our session state variables 

        Args:
            role (str): the role of the message sender
            content (str): the message sent 
        """
        st.session_state.gui_message_history.append({'role': role, 'content': content})
        st.session_state.transcript_history.append({
            'time': datetime.now(timezone.utc).isoformat(timespec='milliseconds'), 
            'session_id': st.session_state.session_id, 
            'user': st.session_state.username, 
            'role': role, 
            'content': content 
        })


    def check_closing_messages(self, msg:str) -> Tuple[bool, str]: 
        """Check if a message contains any of the closing messages 

        Args:
            msg (str): the message to check 

        Returns:
            Tuple[bool, str]: a tuple that returns a bool of whether a closing message was found and a string of the final message 
        """
        for c, m in self.closing_messages.items(): 
            if c.lower() in msg.lower() or m.lower() in msg.lower(): 
                return True, m 
        return False, msg 